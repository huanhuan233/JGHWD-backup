from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import io
import shutil
import re
import tempfile
import subprocess
import platform
from django.http import HttpResponse
from typing import List, Dict, Any, Optional
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK, WD_LINE_SPACING

BASE_CONTENT_PATH = os.path.join(os.path.dirname(__file__))
HEADER_MAPPING = {
    "默认信纸": {
        "type": "image",
        "path": os.path.join(BASE_CONTENT_PATH, "image.png"),
        "width": Inches(6)
    },
    "公司红头": {
        "type": "image",
        "path": os.path.join(BASE_CONTENT_PATH, "company_logo.png"),
        "width": Inches(5)
    },
    "军工信纸": {
        "type": "docx_template",
        "path": os.path.join(BASE_CONTENT_PATH, "contract_header.docx")
    }
}

# 仅保留需求中的字体映射
FONT_MAPPING = {
    "宋体": "SimSun",
    "微软雅黑": "Microsoft YaHei", 
    "黑体": "SimHei",
    "Arial": "Arial",
    "Times New Roman": "Times-Roman"
}

class OptimizedMarkdownParser:
    """优化后的Markdown解析器"""
    
    def __init__(self):
        # 预编译所有正则表达式 - 修复格式冲突
        self.patterns = {
            'header': re.compile(r'^(#{1,6})\s+(.*)$'),
            'bold': re.compile(r'\*\*(.*?)\*\*|__(.*?)__'),
            'italic': re.compile(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)|_(.*?)_'),  # 修复斜体匹配
            'code': re.compile(r'`(.*?)`'),
            'link': re.compile(r'\[(.*?)\]\((.*?)\)'),
            'image': re.compile(r'!\[(.*?)\]\((.*?)\)'),
            'unordered_list': re.compile(r'^[-*+]\s+(.*)$'),
            'ordered_list': re.compile(r'^\d+\.\s+(.*)$'),
            'blockquote': re.compile(r'^>\s+(.*)$'),
            'table_row': re.compile(r'^\|(.+)\|$'),
            'table_separator': re.compile(r'^\|([:\-\s|]+)\|$'),
        }
        
        # 缓存已解析的结果
        self._cache = {}
    
    def parse(self, markdown_text: str) -> List[Dict[str, Any]]:
        """解析Markdown文本 - 优化版本"""
        # 检查缓存
        cache_key = hash(markdown_text)
        if cache_key in self._cache:
            return self._cache[cache_key]
        
        lines = markdown_text.split('\n')
        elements = []  # 直接使用空列表
        i = 0
        total_lines = len(lines)
        
        while i < total_lines:
            line = lines[i]
            
            # 快速检查空行
            if not line.strip():
                elements.append({'type': 'empty'})
                i += 1
                continue
            
            # 使用快速路径检查常见模式
            element = self._quick_parse_line(line, lines, i, total_lines)
            if element:
                elements.append(element['element'])
                i += element['skip_lines']
                continue
            
            # 普通段落（优化处理连续段落）
            paragraph_lines = self._parse_paragraph_fast(lines, i, total_lines)
            if paragraph_lines:
                paragraph_text = ' '.join(paragraph_lines)
                elements.append({
                    'type': 'paragraph', 
                    'text': paragraph_text
                })
                i += len(paragraph_lines)
                continue
            
            i += 1
        
        # 缓存结果
        self._cache[cache_key] = elements
        return elements
        
    def _quick_parse_line(self, line: str, lines: List[str], current_index: int, total_lines: int) -> Optional[Dict[str, Any]]:
        """快速解析单行 - 优化版本"""
        stripped_line = line.strip()
        
        # 快速检查标题（最常见的情况）
        if stripped_line.startswith('#'):
            match = self.patterns['header'].match(stripped_line)
            if match:
                level = len(match.group(1))
                return {
                    'element': {'type': 'header', 'level': level, 'text': match.group(2)},
                    'skip_lines': 1
                }
        
        # 快速检查列表项 - 修复列表识别
        if self._is_list_item(stripped_line):
            list_items = self._parse_list_fast(lines, current_index, total_lines)
            if list_items:
                return {
                    'element': {'type': 'list', 'items': list_items, 'ordered': False},
                    'skip_lines': len(list_items)
                }
        
        # 快速检查有序列表
        if self._is_ordered_list_item(stripped_line):
            ordered_list_items = self._parse_ordered_list_fast(lines, current_index, total_lines)
            if ordered_list_items:
                return {
                    'element': {'type': 'list', 'items': ordered_list_items, 'ordered': True},
                    'skip_lines': len(ordered_list_items)
                }
        
        # 快速检查引用
        if stripped_line.startswith('>'):
            blockquote_lines = self._parse_blockquote_fast(lines, current_index, total_lines)
            if blockquote_lines:
                return {
                    'element': {'type': 'blockquote', 'lines': blockquote_lines},
                    'skip_lines': len(blockquote_lines)
                }
        
        # 快速检查表格
        if stripped_line.startswith('|'):
            table_data = self._parse_table_fast(lines, current_index, total_lines)
            if table_data and len(table_data) > 1:
                return {
                    'element': {'type': 'table', 'data': table_data},
                    'skip_lines': len(table_data) + 1  # 包括分隔行
                }
        
        return None
    
    def _is_list_item(self, line: str) -> bool:
        """判断是否为无序列表项 - 修复识别逻辑"""
        if not line:
            return False
        
        # 检查无序列表标记（-、*、+）
        if len(line) >= 2 and line[0] in ['-', '*', '+'] and line[1] == ' ':
            return True
        
        return False
    
    def _is_ordered_list_item(self, line: str) -> bool:
        """判断是否为有序列表项"""
        if not line:
            return False
        
        # 检查数字列表项（如 "1. "、"2. "等）
        if re.match(r'^\d+\.\s+', line):
            return True
        
        return False
    
    def _parse_table_fast(self, lines: List[str], start_index: int, total_lines: int) -> Optional[List[List[str]]]:
        """快速解析表格 - 简化版本"""
        table_data = []
        i = start_index
        # 解析表头行
        if i >= total_lines:
            return None
            
        header_line = lines[i].strip()
        if not header_line.startswith('|'):
            return None
        
        # 解析表头单元格
        header_cells = [cell.strip() for cell in header_line.split('|')[1:-1]]
        if not header_cells:
            return None
            
        table_data.append(header_cells)
        i += 1
        
        # 检查分隔行（可选）
        if i < total_lines:
            separator_line = lines[i].strip()
            if (separator_line.startswith('|') and 
                any(c in separator_line for c in ['-', ':', '='])):
                i += 1  # 跳过分隔行
        
        # 解析数据行
        while i < total_lines:
            line = lines[i].strip()
            if not line.startswith('|'):
                break
                
            # 解析数据单元格
            data_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(data_cells) == len(header_cells):  # 确保列数匹配
                table_data.append(data_cells)
            i += 1
        
        # 确保至少有两行（表头+数据）
        return table_data if len(table_data) >= 2 else None
    
    def _parse_list_fast(self, lines: List[str], start_index: int, total_lines: int) -> List[str]:
        """快速解析无序列表"""
        items = []
        i = start_index
        
        while i < total_lines:
            line = lines[i].strip()
            
            # 检查是否为无序列表项
            if not self._is_list_item(line):
                break
            
            # 提取内容（优化字符串操作）
            content = line[2:].strip()  # 跳过"- "或"* "等
            
            items.append(content)
            i += 1
        
        return items
    
    def _parse_ordered_list_fast(self, lines: List[str], start_index: int, total_lines: int) -> List[str]:
        """快速解析有序列表"""
        items = []
        i = start_index
        
        while i < total_lines:
            line = lines[i].strip()
            
            # 检查是否为有序列表项
            if not self._is_ordered_list_item(line):
                break
            
            # 提取内容
            content = re.sub(r'^\d+\.\s+', '', line)  # 移除数字标记
            
            items.append(content)
            i += 1
        
        return items
    
    def _parse_blockquote_fast(self, lines: List[str], start_index: int, total_lines: int) -> List[str]:
        """快速解析引用块"""
        blockquote_lines = []
        i = start_index
        
        while i < total_lines:
            line = lines[i].strip()
            if not line.startswith('>'):
                break
            
            # 快速提取内容
            content = line[1:].strip()  # 跳过">"
            if content.startswith(' '):  # 跳过可能存在的空格
                content = content[1:]
            
            blockquote_lines.append(content)
            i += 1
        
        return blockquote_lines
    
    def _parse_paragraph_fast(self, lines: List[str], start_index: int, total_lines: int) -> List[str]:
        """快速解析段落 - 优化连续段落处理"""
        paragraph_lines = []
        i = start_index
        
        while i < total_lines:
            line = lines[i]
            stripped_line = line.strip()
            
            # 快速检查是否为新元素开始
            if not stripped_line:
                break
            
            # 使用第一个字符快速判断
            first_char = stripped_line[0] if stripped_line else ''
            if first_char in ['#', '>', '|', '-', '*', '+']:
                # 进一步验证是否真的是特殊元素
                if (first_char == '#' and self.patterns['header'].match(stripped_line)) or \
                   (first_char == '>' and self.patterns['blockquote'].match(stripped_line)) or \
                   (first_char == '|' and self.patterns['table_row'].match(stripped_line)) or \
                   (self._is_list_item(stripped_line) or self._is_ordered_list_item(stripped_line)):
                    break
            
            paragraph_lines.append(stripped_line)
            i += 1
        
        return paragraph_lines
    
    def parse_inline_formatting(self, text: str) -> str:
        """解析行内格式 - 优化版本，修复格式冲突"""
        if not text:
            return text
        
        # 缓存结果
        cache_key = f"inline_{hash(text)}"
        if cache_key in self._cache:
            return self._cache[cache_key]
        
        # 使用单个替换操作减少字符串操作次数
        def replace_bold(match):
            content = match.group(1) or match.group(2)
            return f'<bold>{content}</bold>'
        
        def replace_italic(match):
            # 修复斜体匹配，避免与列表标记冲突
            content = match.group(1) or match.group(2)
            return f'<italic>{content}</italic>'
        
        def replace_code(match):
            return f'<code>{match.group(1)}</code>'
        
        def replace_link(match):
            return f'<link text="{match.group(1)}" url="{match.group(2)}">'
        
        # 批量替换（减少字符串复制）
        result = text
        result = self.patterns['bold'].sub(replace_bold, result)
        result = self.patterns['italic'].sub(replace_italic, result)
        result = self.patterns['code'].sub(replace_code, result)
        result = self.patterns['link'].sub(replace_link, result)
        
        # 缓存结果
        self._cache[cache_key] = result
        return result

    def clear_cache(self):
        """清空缓存"""
        self._cache.clear()

class LibreOfficePDFConverter:
    """LibreOffice PDF转换器"""
    
    def __init__(self):
        self.libreoffice_path = self._find_libreoffice()
    
    def _find_libreoffice(self):
        """查找LibreOffice安装路径"""
        system = platform.system()
        
        if system == "Windows":
            # Windows常见安装路径
            possible_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"D:\Program Files\LibreOffice\program\soffice.exe",
                r"D:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"F:\Program Files\LibreOffice\program\soffice.exe",
            ]
        elif system == "Darwin":  # macOS
            possible_paths = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            ]
        else:  # Linux
            possible_paths = [
                "/usr/bin/libreoffice",
                "/usr/bin/soffice",
                "/snap/bin/libreoffice",
                "/opt/libreoffice/program/soffice",
            ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        # 如果在标准路径找不到，尝试在PATH中查找
        soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice_path:
            return soffice_path
        
        raise Exception("未找到LibreOffice安装路径，请确保已安装LibreOffice")

    def convert_to_pdf(self, docx_stream, output_pdf_path=None):
        """生成PDF并返回路径，暂不清理临时文件"""
        temp_dir = tempfile.mkdtemp()
        docx_temp_path = os.path.join(temp_dir, "temp_document.docx")
        
        try:
            # 保存DOCX到临时文件
            docx_stream.seek(0)
            with open(docx_temp_path, "wb") as f:
                f.write(docx_stream.getvalue())
            
            # 构建转换命令
            cmd = [
                self.libreoffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", temp_dir,
                docx_temp_path
            ]
            
            # 执行转换
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                universal_newlines=True,
                timeout=60
            )
            
            if result.returncode != 0:
                error_msg = f"转换失败: {result.stderr}"
                raise Exception(error_msg)
            
            # 查找生成的PDF文件
            expected_pdf_path = os.path.join(temp_dir, "temp_document.pdf")
            if not os.path.exists(expected_pdf_path):
                pdf_files = [f for f in os.listdir(temp_dir) if f.endswith('.pdf')]
                if not pdf_files:
                    raise Exception("PDF文件未生成")
                expected_pdf_path = os.path.join(temp_dir, pdf_files[0])
            
            # 如果指定了输出路径则复制，否则返回临时路径
            if output_pdf_path:
                shutil.copy2(expected_pdf_path, output_pdf_path)
                # 复制后清理临时目录（因为输出路径已保存）
                shutil.rmtree(temp_dir, ignore_errors=True)
                return output_pdf_path
            else:
                # 未指定输出路径时，返回临时路径并保留目录（由调用方清理）
                return expected_pdf_path
                
        except subprocess.TimeoutExpired:
            raise Exception("转换超时")
        except Exception as e:
            # 发生错误时清理临时目录
            shutil.rmtree(temp_dir, ignore_errors=True)
            raise Exception(f"PDF转换错误: {str(e)}")
        # 移除finally块的强制清理，改为按需清理
    
    def convert_to_pdf_stream(self, docx_stream):
        """读取PDF流后再清理临时文件"""
        temp_pdf_path = None
        temp_dir = None
        try:
            temp_pdf_path = self.convert_to_pdf(docx_stream)
            if not temp_pdf_path or not os.path.exists(temp_pdf_path):
                raise Exception("PDF文件生成失败")
            
            # 获取临时目录路径（用于后续清理）
            temp_dir = os.path.dirname(temp_pdf_path)
            
            # 读取PDF内容到流
            with open(temp_pdf_path, "rb") as f:
                pdf_stream = io.BytesIO(f.read())
            
            return pdf_stream
            
        except Exception as e:
            raise Exception(f"读取PDF流失败: {str(e)}")
        finally:
            # 确保临时文件和目录被清理
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)

class DocGenerator:
    def __init__(self, content_format="docx"):
        self.document = Document()
        self.content_format = content_format.lower()
        self.elements = []
        self.markdown_parser = OptimizedMarkdownParser()
        self.pdf_converter = LibreOfficePDFConverter() if content_format == "pdf" else None
        self.current_header = {
            "type": None,
            "path": None,
            "width": Inches(6)
        }

        # 新增：记录当前标头类型（用于预览和替换）
        self.current_header_type = None  # "image" / "docx_template" / None
        # 新增：记录当前标头源（文件路径或模板内容，用于替换）
        self.current_header_source = None

    def _get_header_section(self):
        """获取 Word 文档的"首要节"（确保标头应用到整个文档）"""
        if not self.document.sections:
            return self.document.sections[0]
        
        for section in self.document.sections:
            section.different_first_page_header_footer = False
            section.header_distance = Pt(15)
        return self.document.sections[0]

    
    def clear_header(self):
        for section in self.document.sections:
            for para in section.header.paragraphs:
                para.clear()
            for table in section.header.tables:
                table._element.getparent().remove(table._element)
        self.current_header = {"type": None, "path": None, "width": Inches(6)}

    def preview_header(self):
        if not self.current_header["type"]:
            return {"status": "empty", "message": "当前无标头"}
        
        section = self._get_header_section()
        header_paragraphs = section.header.paragraphs
        return {
            "type": self.current_header["type"],
            "path": self.current_header["path"],
            "paragraph_count": len(header_paragraphs),
            "content": [{
                "text": para.text,
                "alignment": para.alignment.name,
                "has_image": bool(para.runs and para.runs[0].element.xpath('.//a:blip'))
            } for para in header_paragraphs]
        }

    def set_header_alignment(self, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
        """设置标头的整体对齐方式"""
        if self.current_header_type is None:
            raise ValueError("当前无标头，无法设置对齐方式")
        
        section = self._get_header_section()
        for para in section.header.paragraphs:
            para.alignment = alignment
        return True

    def set_header_margin(self, distance=Pt(15)):
        """设置标头与页边距的距离"""
        for section in self.document.sections:
            section.header_distance = distance
        return True

    def load_markdown(self, markdown_text):
        """加载并解析Markdown文本"""
        if len(markdown_text) > 100000:
            self._load_large_markdown(markdown_text)
        else:
            self.markdown_elements = self.markdown_parser.parse(markdown_text)
            self._convert_markdown_to_docx()

    def _load_large_markdown(self, markdown_text):
        """处理大型Markdown文件（分块处理）"""
        chunks = markdown_text.split('\n\n')
        self.markdown_elements = []
        
        for i, chunk in enumerate(chunks):
            if i % 100 == 0:
                self.markdown_parser.clear_cache()
            
            if chunk.strip():
                elements = self.markdown_parser.parse(chunk)
                self.markdown_elements.extend(elements)
        
        self._convert_markdown_to_docx()

    def _convert_markdown_to_docx(self):
        """将Markdown元素转换为Word文档内容"""
        for element in self.markdown_elements:
            if element['type'] == 'header':
                self.add_heading(element['text'], element['level'])
            elif element['type'] == 'paragraph':
                self._add_formatted_paragraph(element['text'])
            elif element['type'] == 'table':
                self._add_markdown_table(element['data'])
            elif element['type'] == 'list':
                self._add_markdown_list(element['items'], element.get('ordered', False))
            elif element['type'] == 'blockquote':
                self._add_markdown_blockquote(element['lines'])
            elif element['type'] == 'empty':
                self.document.add_paragraph()

    def _add_formatted_paragraph(self, text, target_paragraph=None):
        """添加带格式的段落"""
        formatted_text = self.markdown_parser.parse_inline_formatting(text)
        paragraph = target_paragraph if target_paragraph is not None else self.document.add_paragraph()
        
        parts = re.split(r'(<[^>]+>)', formatted_text)
        current_text = ""
        bold = False
        italic = False
        code = False
        
        for part in parts:
            if part == '<bold>':
                bold = True
            elif part == '</bold>':
                bold = False
            elif part == '<italic>':
                italic = True
            elif part == '</italic>':
                italic = False
            elif part.startswith('<code>'):
                code = True
                current_text = part.replace('<code>', '')
            elif part == '</code>':
                self._add_run(paragraph, current_text, bold, italic, code, is_code_end=True)
                code = False
                current_text = ""
            elif part.startswith('<link'):
                match = re.search(r'text="(.*?)" url="(.*?)"', part)
                if match:
                    link_text = match.group(1)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
            elif part.startswith('<'):
                continue
            else:
                current_text += part
                if not part.startswith('<'):
                    self._add_run(paragraph, current_text, bold, italic, code)
                    current_text = ""
        
        if current_text:
            self._add_run(paragraph, current_text, bold, italic, code)
        
        return paragraph

    def _add_run(self, paragraph, text, bold=False, italic=False, code=False, is_code_end=False):
        """添加文本运行"""
        if not text.strip():
            return
        
        run = paragraph.add_run(text)
        
        if code or is_code_end:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 0, 0)
        else:
            run.font.name = "宋体"
            run.font.size = Pt(12)
            if os.name == "posix":
                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), "SimSun")

        run.font.bold = bold
        run.font.italic = italic

    def _add_markdown_table(self, table_data):
        """添加Markdown表格"""
        if not table_data or len(table_data) < 2:
            return
        
        rows = len(table_data)
        cols = len(table_data[0])
        
        table = self.add_table(rows, cols, style="default")
        
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = ''
                cell_paragraph = cell.paragraphs[0]
                self._add_formatted_paragraph(cell_text, target_paragraph=cell_paragraph)
                self._set_cell_padding(cell, padding=5)

    def _set_cell_padding(self, cell, padding=5):
        """设置表格单元格内边距"""
        tc_pr = cell._element.get_or_add_tcPr()
        
        sides = {
            'top': 'top',
            'bottom': 'bottom',
            'left': 'left',
            'right': 'right'
        }
        
        for side in sides:
            mar_element = OxmlElement(f'w:tc{side.capitalize()}')
            w_element = OxmlElement('w:w')
            w_element.set(qn('w:val'), str(padding * 20))
            w_element.set(qn('w:type'), 'dxa')
            mar_element.append(w_element)
            tc_pr.append(mar_element)

    def _add_markdown_list(self, items, ordered=False):
        """添加Markdown列表 - 修复列表处理"""
        for index, item in enumerate(items):
            paragraph = self.document.add_paragraph()
            paragraph.style = self.document.styles['List Paragraph']
            
            # 添加列表标记
            if ordered:
                marker = f'{index + 1}. '
            else:
                marker = '• '
            
            run = paragraph.add_run(marker)
            run.font.name = "宋体"
            
            # 添加列表内容（支持内联格式）
            self._add_formatted_paragraph(item, target_paragraph=paragraph)

    def _add_markdown_blockquote(self, lines):
        """添加Markdown引用块"""
        for line in lines:
            paragraph = self.document.add_paragraph()
            paragraph.style = self.document.styles['Intense Quote']
            run = paragraph.add_run(line)
            run.font.name = "宋体"
            run.font.italic = True
            if os.name == "posix":
                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), "SimSun")

    def set_text_style(self, font_name="宋体", font_size=14, bold=False, italic=False):
        """设置文本样式"""
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                self._set_run_style(run, font_name, font_size, bold, italic)
        
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._set_run_style(run, font_name, font_size, bold, italic)

    def _set_run_style(self, run, font_name, font_size, bold, italic, color=None):
        """修复字体设置方法"""
        # 使用字体映射
        mapped_font = FONT_MAPPING.get(font_name, "SimSun")
        
        # 设置英文字体
        run.font.name = mapped_font
        run._element.rPr.rFonts.set(qn('w:ascii'), mapped_font)
        
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), mapped_font)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), mapped_font)
        
        # 强制设置加粗
        if bold:
            run.font.bold = True
        else:
            run.font.bold = False  # 明确设置为False，避免继承问题
        
        run.font.size = Pt(font_size)
        run.font.italic = italic
        
        # 设置颜色（默认为黑色）
        if color is not None:
            # font_color是RGBColor对象（如RGBColor(255,0,0)）
            run.font.color.rgb = color
        else:
            run.font.color.rgb = RGBColor(0, 0, 0)  # 默认黑色
        
        # 确保无下划线
        run.font.underline = None

    def add_heading(self, text, level=1):
        """修复标题加粗问题 - 重新实现确保加粗生效"""
        # 直接创建段落并手动设置样式，避免默认样式覆盖
        paragraph = self.document.add_paragraph()
        
        # 设置段落样式为标题
        if level == 1:
            paragraph.style = self.document.styles['Heading 1']
        elif level == 2:
            paragraph.style = self.document.styles['Heading 2']
        elif level == 3:
            paragraph.style = self.document.styles['Heading 3']
        else:
            paragraph.style = self.document.styles['Heading 4']
        
        # 修改样式定义本身，确保样式不包含keepNext和keepLines
        try:
            style = paragraph.style
            style_element = style.element
            ppr_style = style_element.find(qn('w:pPr'))
            if ppr_style is not None:
                # 移除样式中的keepNext
                keep_next_style = ppr_style.find(qn('w:keepNext'))
                if keep_next_style is not None:
                    ppr_style.remove(keep_next_style)
                # 移除样式中的keepLines
                keep_lines_style = ppr_style.find(qn('w:keepLines'))
                if keep_lines_style is not None:
                    ppr_style.remove(keep_lines_style)
        except Exception as e:
            print(f"修改样式定义失败: {e}")
            # 继续执行，段落级别的设置仍然有效
        
        # 清除现有runs
        for run in paragraph.runs:
            run.clear()
        
        # 添加新run
        run = paragraph.add_run(text)
        
        # 设置字体
        font_name = "黑体" if level == 1 else "宋体"
        mapped_font = FONT_MAPPING.get(font_name, "SimHei" if level == 1 else "SimSun")
        
        # 强制设置字体
        run.font.name = mapped_font
        run._element.rPr.rFonts.set(qn('w:ascii'), mapped_font)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), mapped_font)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), mapped_font)
        
        # 强制加粗 - 使用XML级别设置确保生效
        run.font.bold = False
        # 额外通过XML设置确保加粗
        if level != 1:  # 重点：一级标题不执行这段XML加粗逻辑
            rpr = run._element.get_or_add_rPr()
            b = OxmlElement('w:b')
            b.set(qn('w:val'), 'true')
            rpr.append(b)
        
        # 设置字体大小
        font_sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10}
        run.font.size = Pt(font_sizes.get(level, 12))
        
        # 设置颜色为黑色
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 设置段落对齐
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # 设置段落格式：段前0磅，取消"与下段同页"和"段中不分页"
        para_format = paragraph.paragraph_format
        para_format.space_before = Pt(0)  # 段前0磅
        para_format.space_after = Pt(0)   # 段后0磅（可选）
        
        # 获取段落属性元素
        ppr = paragraph._element.get_or_add_pPr()
        
        # 强制取消"与下段同页"（keepNext）- 先删除再显式设置为false
        keep_next = ppr.find(qn('w:keepNext'))
        if keep_next is not None:
            ppr.remove(keep_next)
        # 显式设置为false
        keep_next_elem = OxmlElement('w:keepNext')
        keep_next_elem.set(qn('w:val'), 'false')
        ppr.append(keep_next_elem)
        
        # 强制取消"段中不分页"（keepLines）- 先删除再显式设置为false
        keep_lines = ppr.find(qn('w:keepLines'))
        if keep_lines is not None:
            ppr.remove(keep_lines)
        # 显式设置为false
        keep_lines_elem = OxmlElement('w:keepLines')
        keep_lines_elem.set(qn('w:val'), 'false')
        ppr.append(keep_lines_elem)
        
        # 确保widowControl设置为false（段中不分页相关）
        widow_control = ppr.find(qn('w:widowControl'))
        if widow_control is not None:
            widow_control.set(qn('w:val'), 'false')
        else:
            widow_control_elem = OxmlElement('w:widowControl')
            widow_control_elem.set(qn('w:val'), 'false')
            ppr.append(widow_control_elem)
        
        return paragraph

    def add_table(self, rows, cols, style="default"):
        """添加表格"""
        table = self.document.add_table(rows=rows, cols=cols)
        if style == "no_border":
            for row in table.rows:
                for cell in row.cells:
                    for border in cell._element.xpath(".//w:tcBorders/*"):
                        border.set("w:val", "none")
        elif style == "striped":
            table.style = self.document.styles["Light Shading"]
        else:
            table.style = self.document.styles["Table Grid"]
        return table

    def add_content(self, text, font_name="宋体", font_size=12, bold=False, italic=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
        """添加内容"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        self._set_run_style(run, font_name, font_size, bold, italic)
        paragraph.alignment = alignment
        return paragraph

    def add_page_break(self):
        """添加分页符"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
        return paragraph

    def set_table_style(self, style="default"):
        """设置表格样式"""
        self.table_style = style
        for table in self.document.tables:
            self._apply_table_style(table, style)

    def _apply_table_style(self, table, style):
        """应用具体表格样式"""
        tbl_pr = table._tblPr
        
        existing_borders = tbl_pr.xpath('.//w:tblBorders')
        for border in existing_borders:
            border.getparent().remove(border)
        
        existing_shading = tbl_pr.xpath('.//w:shd')
        for shading in existing_shading:
            shading.getparent().remove(shading)
        
        if style == "no_border":
            tbl_borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tbl_borders.append(border)
            tbl_pr.append(tbl_borders)
            
        elif style == "striped":
            for i, row in enumerate(table.rows):
                if i % 2 == 1:
                    for cell in row.cells:
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'F2F2F2')
                        cell._element.get_or_add_tcPr().append(shading)
            
            tbl_borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tbl_borders.append(border)
            tbl_pr.append(tbl_borders)
            
        else:
            tbl_borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tbl_borders.append(border)
            tbl_pr.append(tbl_borders)

    def save(self):
        """保存文档"""
        try:
            if self.content_format == "docx":
                file_stream = io.BytesIO()
                self.document.save(file_stream)
                file_stream.seek(0)
                response = HttpResponse(
                    file_stream.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = 'attachment; filename="generated_document.docx"'
                response['Content-Length'] = file_stream.getbuffer().nbytes
                return response
                
            elif self.content_format == "pdf":
                # 先保存为DOCX流
                docx_stream = io.BytesIO()
                self.document.save(docx_stream)
                docx_stream.seek(0)  # 确保流指针在起始位置
                
                # 使用LibreOffice转换为PDF
                pdf_stream = self.pdf_converter.convert_to_pdf_stream(docx_stream)
                
                response = HttpResponse(
                    pdf_stream.getvalue(),
                    content_type='application/pdf'
                )
                response['Content-Disposition'] = 'attachment; filename="generated_document.pdf"'
                response['Content-Length'] = pdf_stream.getbuffer().nbytes
                return response
                
            else:
                return HttpResponse("不支持的文件格式", status=400)
                
        except Exception as e:
            error_msg = f"保存文档出错：{str(e)}"
            print(error_msg)
            return HttpResponse(error_msg, status=500)
    def set_header_by_frontend_param(self, frontend_param: str) -> bool:
            if frontend_param not in HEADER_MAPPING:
                raise ValueError(f"前端标头参数 {frontend_param} 未配置")
            
            header_config = HEADER_MAPPING[frontend_param]
            self.current_header["type"] = header_config["type"]
            self.current_header["path"] = header_config["path"]
            if "width" in header_config:
                self.current_header["width"] = header_config["width"]
            
            if not os.path.exists(header_config["path"]):
                raise FileNotFoundError(f"标头文件不存在：{header_config['path']}")
            
            if header_config["type"] == "image":
                ext = os.path.splitext(header_config["path"])[-1].lower()
                if ext not in [".png", ".jpg", ".jpeg"]:
                    raise ValueError(f"图片标头仅支持PNG/JPG，当前格式：{ext}")
                self._add_image_header(header_config["path"], header_config["width"])
            
            elif header_config["type"] == "docx_template":
                if not header_config["path"].lower().endswith(".docx"):
                    raise ValueError("DOCX模板标头必须是.docx格式")
                self._add_docx_template_header(header_config["path"])
            
            return True

    def _add_docx_template_header(self, docx_path: str):
        self.clear_header()
        template_doc = Document(docx_path)
        if not template_doc.sections or not template_doc.sections[0].header.paragraphs:
            raise ValueError("DOCX模板中未找到有效标头")
        
        template_header = template_doc.sections[0].header
        current_header = self._get_header_section().header
        
        for template_para in template_header.paragraphs:
            new_para = current_header.add_paragraph()
            new_para.text = template_para.text
            new_para.style = template_para.style
            
            for run in template_para.runs:
                if run.element.xpath('.//a:blip'):
                    for blip in run.element.xpath('.//a:blip'):
                        rId = blip.get(qn('r:embed'))
                        if rId in template_doc.part.related_parts:
                            img_part = template_doc.part.related_parts[rId]
                            img_stream = io.BytesIO(img_part._blob)
                            new_run = new_para.add_run()
                            new_run.add_picture(img_stream, width=run.element.xpath('.//a:ext')[0].get(qn('cx')))
    def _add_image_header(self, img_path: str, width: Inches):
        self.clear_header()
        section = self._get_header_section()
        header_para = section.header.add_paragraph()
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run = header_para.add_run()
        run.add_picture(img_path, width=width)

    def set_line_spacing(self, line_spacing_type, line_spacing_value=None):
        """设置文档的行间距
        :param line_spacing_type: 行间距类型，可选值：'single'（单倍）、'one_point_five'（1.5倍）、'double'（2倍）、'at_least'（最小值）、'exact'（固定值）、'multiple'（多倍行距）
        :param line_spacing_value: 当行间距类型为'exact'或'multiple'时，指定具体的磅数，如20、25等
        """
        line_spacing_map = {
            'single': WD_LINE_SPACING.SINGLE,
            'one_point_five': WD_LINE_SPACING.ONE_POINT_FIVE,
            'double': WD_LINE_SPACING.DOUBLE,
            'at_least': WD_LINE_SPACING.AT_LEAST,
            'exact': WD_LINE_SPACING.EXACTLY,
            'multiple': WD_LINE_SPACING.MULTIPLE
        }
        if line_spacing_type not in line_spacing_map:
            raise ValueError(f"不支持的行间距类型：{line_spacing_type}，可选类型：{list(line_spacing_map.keys())}")
        
        # 对所有段落应用行间距
        for paragraph in self.document.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing_rule = line_spacing_map[line_spacing_type]
            # 如果是"固定值"或"多倍行距"，且指定了数值，则设置具体磅数
            if line_spacing_type in ['exact', 'multiple'] and line_spacing_value is not None:
                paragraph_format.line_spacing = Pt(line_spacing_value)
    def get_all_heading(self) -> List[Dict[str, Any]]:
            """
            获取文档中所有一级标题（Heading 1）
            返回格式：[{'index': 段落索引, 'text': 标题文本}, ...]
            """
            heading1_list = []
            # 遍历所有段落，筛选样式为"Heading 1"的段落
            for para_index, paragraph in enumerate(self.document.paragraphs):
                # 判断段落样式是否为一级标题（兼容中英文环境）
                if paragraph.style.name in ["Heading 1", "标题 1"]:
                    # 提取段落纯文本（排除格式标记）
                    heading_text = paragraph.text.strip()
                    if heading_text:  # 过滤空标题
                        heading1_list.append({
                            "index": para_index,  # 段落在文档中的索引（便于定位）
                            "text": heading_text   # 一级标题文本
                        })
            return heading1_list
    def match_headings(self, title_setting, all_headings):
        """匹配并设置标题样式"""
        # 首先需要建立段落索引映射
        paragraph_map = {}
        for i, para in enumerate(self.document.paragraphs):
            paragraph_map[i] = para
        
        for item in title_setting.get('structure', []):
            # 匹配一级标题
            for heading in all_headings:
                if heading["text"] == item["title"]:
                    print(f"找到匹配的一级标题: {item['title']}")
                    
                    # 获取对应的段落对象
                    para_index = heading["index"]
                    paragraph = paragraph_map.get(para_index)
                    
                    if paragraph:
                        # 处理匹配的一级标题
                        current_font = item["font"]
                        current_size = item["size"]
                        current_bold = item["bold"]
                        current_italic = item["italic"]
                        color_str = item["color"]
                        
                        # 解析颜色
                        color_numbers = re.findall(r'\d+', color_str)
                        if len(color_numbers) >= 3:
                            r, g, b = map(int, color_numbers[:3])
                            font_color = RGBColor(r, g, b)
                        else:
                            font_color = RGBColor(0, 0, 0)  # 默认黑色
                        
                        # 设置段落中所有run的样式
                        for run in paragraph.runs:
                            self._set_run_style(
                                run,  # 位置参数
                                current_font,
                                current_size,
                                current_bold,
                                current_italic,
                                font_color  
                            )
                        
                        # 设置段落格式：段前0磅，取消"与下段同页"和"段中不分页"
                        para_format = paragraph.paragraph_format
                        para_format.space_before = Pt(0)  # 段前0磅
                        
                        # 获取段落属性元素
                        ppr = paragraph._element.get_or_add_pPr()
                        
                        # 强制取消"与下段同页"（keepNext）
                        keep_next = ppr.find(qn('w:keepNext'))
                        if keep_next is not None:
                            ppr.remove(keep_next)
                        # 显式设置为false
                        keep_next_elem = OxmlElement('w:keepNext')
                        keep_next_elem.set(qn('w:val'), 'false')
                        ppr.append(keep_next_elem)
                        
                        # 强制取消"段中不分页"（keepLines）
                        keep_lines = ppr.find(qn('w:keepLines'))
                        if keep_lines is not None:
                            ppr.remove(keep_lines)
                        # 显式设置为false
                        keep_lines_elem = OxmlElement('w:keepLines')
                        keep_lines_elem.set(qn('w:val'), 'false')
                        ppr.append(keep_lines_elem)
                        
                        # 确保widowControl设置为false
                        widow_control = ppr.find(qn('w:widowControl'))
                        if widow_control is not None:
                            widow_control.set(qn('w:val'), 'false')
                        else:
                            widow_control_elem = OxmlElement('w:widowControl')
                            widow_control_elem.set(qn('w:val'), 'false')
                            ppr.append(widow_control_elem)
            
            
            if 'children' in item and item['children']:
                self.match_headings({'structure': item['children']}, all_headings)
    
    def add_table_of_contents(self, max_level=2, font_name="宋体"):
        """
        在文档开头添加目录（直接生成目录内容）
        :param max_level: 目录最大层级（1-10）
        :param font_name: 目录字体名称，与正文字体相同
        """
        # 确保 max_level 在有效范围内
        max_level = max(1, min(10, max_level))
        
        # 获取所有标题
        headings = self._get_all_headings_with_levels(max_level)
        
        if not headings:
            # 如果没有标题，添加一个提示
            if len(self.document.paragraphs) == 0:
                self.document.add_paragraph()
            first_paragraph = self.document.paragraphs[0]
            toc_paragraph = first_paragraph.insert_paragraph_before()
            run = toc_paragraph.add_run("目录")
            self._set_run_style(run, font_name, 16, True, False)
            toc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.document.paragraphs[0].insert_paragraph_before()
            return
        
        # 在文档开头插入目录
        if len(self.document.paragraphs) == 0:
            self.document.add_paragraph()
        
        first_paragraph = self.document.paragraphs[0]
        
        # 添加"目录"标题
        toc_title_para = first_paragraph.insert_paragraph_before()
        toc_title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = toc_title_para.add_run("目录")
        self._set_run_style(run, font_name, 16, True, False)
        
        # 添加空行
        first_paragraph.insert_paragraph_before()
        
        # 生成目录内容
        for heading in headings:
            level = heading['level']
            text = heading['text']
            para_index = heading['index']
            target_para = heading.get('paragraph')
            
            # 创建目录项段落
            toc_item_para = first_paragraph.insert_paragraph_before()
            
            # 根据层级设置缩进
            indent_size = (level - 1) * 0.5  # 每级缩进0.5英寸
            if indent_size > 0:
                toc_item_para.paragraph_format.left_indent = Inches(indent_size)
            
            # 添加目录项文本（黑色字体，无下划线）
            run = toc_item_para.add_run(text)
            self._set_run_style(run, font_name, 12, False, False, RGBColor(0, 0, 0))  # 黑色
            
            # 添加制表符
            toc_item_para.add_run("\t")
            
            # 为标题段落添加书签，以便PAGEREF字段引用
            if target_para:
                try:
                    # 创建书签名称（基于标题文本，清理特殊字符）
                    bookmark_name = f"_Toc{para_index}_{level}"
                    # 在目标标题段落添加书签
                    self._add_bookmark_to_paragraph(target_para, bookmark_name)
                    
                    # 添加PAGEREF字段来引用标题页码
                    page_run = toc_item_para.add_run()
                    self._add_pageref_field(page_run, bookmark_name)
                    
                    # 设置页码字体为黑色
                    self._set_run_style(page_run, font_name, 12, False, False, RGBColor(0, 0, 0))
                except Exception as e:
                    # 如果添加书签失败，使用占位符
                    print(f"添加页码字段失败: {e}")
                    page_run = toc_item_para.add_run("...")
                    self._set_run_style(page_run, font_name, 12, False, False, RGBColor(0, 0, 0))
            else:
                # 如果没有目标段落，使用占位符
                page_run = toc_item_para.add_run("...")
                self._set_run_style(page_run, font_name, 12, False, False, RGBColor(0, 0, 0))
            
            # 设置制表符位置（右对齐页码）
            ppr = toc_item_para._element.get_or_add_pPr()
            tabs_elem = ppr.find(qn('w:tabs'))
            if tabs_elem is None:
                tabs_elem = OxmlElement('w:tabs')
                ppr.append(tabs_elem)
            
            # 添加右对齐制表符（带点线引导符）
            tab_elem = OxmlElement('w:tab')
            tab_elem.set(qn('w:val'), 'right')
            tab_elem.set(qn('w:leader'), 'dot')
            # 设置制表符位置（约6.5英寸，转换为twips单位，1英寸=1440 twips）
            tab_pos = int(6.5 * 1440)
            tab_elem.set(qn('w:pos'), str(tab_pos))
            tabs_elem.append(tab_elem)
        
        # 添加空行分隔目录和正文
        first_paragraph.insert_paragraph_before()
        
        # 同时添加TOC字段，以便用户按F9可以更新
        self._add_toc_field(max_level, font_name)
    
    def _get_all_headings_with_levels(self, max_level):
        """
        获取文档中所有指定层级范围内的标题
        :param max_level: 最大层级
        :return: 标题列表，包含level、text、index
        """
        headings = []
        style_level_map = {
            'Heading 1': 1, '标题 1': 1,
            'Heading 2': 2, '标题 2': 2,
            'Heading 3': 3, '标题 3': 3,
            'Heading 4': 4, '标题 4': 4,
            'Heading 5': 5, '标题 5': 5,
            'Heading 6': 6, '标题 6': 6,
            'Heading 7': 7, '标题 7': 7,
            'Heading 8': 8, '标题 8': 8,
            'Heading 9': 9, '标题 9': 9,
            'Heading 10': 10, '标题 10': 10,
        }
        
        # 遍历所有段落，收集标题信息
        # 注意：这里收集的是当前文档中的标题，目录会在开头插入
        for para_index, paragraph in enumerate(self.document.paragraphs):
            style_name = paragraph.style.name
            if style_name in style_level_map:
                level = style_level_map[style_name]
                if level <= max_level:
                    text = paragraph.text.strip()
                    if text:
                        headings.append({
                            'level': level,
                            'text': text,
                            'index': para_index,
                            'paragraph': paragraph  # 保存段落引用，用于创建超链接
                        })
        
        return headings
    
    def _add_toc_field(self, max_level, font_name):
        """
        添加TOC字段，以便用户按F9可以更新目录
        """
        try:
            # 找到目录区域的最后一个段落（在正文之前）
            # 在最后一个目录项后添加TOC字段
            if len(self.document.paragraphs) < 2:
                return
            
            # 在目录区域末尾添加TOC字段段落（隐藏，用于更新）
            # 找到正文开始的位置（第一个非目录段落）
            toc_end_index = 0
            for i, para in enumerate(self.document.paragraphs):
                if para.text.strip() and not para.text.strip().startswith('...'):
                    # 检查是否是标题样式
                    if para.style.name not in ['Heading 1', '标题 1', 'Heading 2', '标题 2', 
                                               'Heading 3', '标题 3', 'Heading 4', '标题 4',
                                               'Heading 5', '标题 5', 'Heading 6', '标题 6']:
                        toc_end_index = i
                        break
            
            if toc_end_index > 0:
                # 在目录结束位置插入TOC字段
                toc_para = self.document.paragraphs[toc_end_index - 1]
                toc_field_para = toc_para.insert_paragraph_after()
                
                # 创建标准TOC字段
                toc_instruction = f'TOC \\o "1-{max_level}" \\h \\z \\u'
                
                # 使用完整的字段结构（fldChar + instrText + fldChar）
                # 开始字段字符
                fld_char_begin = OxmlElement('w:fldChar')
                fld_char_begin.set(qn('w:fldCharType'), 'begin')
                
                # 字段指令
                instr_text = OxmlElement('w:instrText')
                instr_text.set(qn('xml:space'), 'preserve')
                instr_text.text = toc_instruction
                
                # 结束字段字符
                fld_char_end = OxmlElement('w:fldChar')
                fld_char_end.set(qn('w:fldCharType'), 'end')
                
                # 创建run并添加字段元素
                run1 = toc_field_para.add_run()
                run1._element.append(fld_char_begin)
                
                run2 = toc_field_para.add_run()
                run2._element.append(instr_text)
                
                run3 = toc_field_para.add_run()
                run3._element.append(fld_char_end)
                
                # 隐藏这个段落（用于更新，不显示）
                ppr = toc_field_para._element.get_or_add_pPr()
                ppr.set(qn('w:vanish'), 'true')
        except Exception as e:
            print(f"添加TOC字段失败: {e}")
            # 不抛出异常，目录内容已经生成
    
    def _set_toc_style(self, font_name="宋体"):
        """
        设置TOC样式以使用指定的字体
        注意：Word中的TOC样式（TOC 1, TOC 2等）控制目录的显示
        我们通过修改样式表来设置这些样式使用正文字体
        """
        try:
            # 获取字体映射
            mapped_font = FONT_MAPPING.get(font_name, "SimSun")
            
            # TOC样式名称列表（TOC 1到TOC 9对应1-9级标题）
            toc_style_names = [f"TOC {i}" for i in range(1, 10)]
            
            # 遍历所有样式，找到TOC样式并设置字体
            for style in self.document.styles:
                if style.name in toc_style_names:
                    try:
                        # 获取样式的XML元素
                        style_element = style.element
                        
                        # 使用XPath查找rPr元素（运行属性）
                        # python-docx使用特定的命名空间
                        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        rpr_list = style_element.xpath('.//w:rPr', namespaces=nsmap)
                        
                        if not rpr_list:
                            # 如果没有rPr，需要创建
                            # 先查找或创建pPr（段落属性）
                            ppr_list = style_element.xpath('.//w:pPr', namespaces=nsmap)
                            if not ppr_list:
                                ppr = OxmlElement('w:pPr')
                                style_element.append(ppr)
                            else:
                                ppr = ppr_list[0]
                            
                            rpr = OxmlElement('w:rPr')
                            ppr.append(rpr)
                        else:
                            rpr = rpr_list[0]
                        
                        # 设置字体
                        rfonts_list = rpr.xpath('.//w:rFonts', namespaces=nsmap)
                        if not rfonts_list:
                            rfonts = OxmlElement('w:rFonts')
                            rfonts.set(qn('w:ascii'), mapped_font)
                            rfonts.set(qn('w:eastAsia'), mapped_font)
                            rfonts.set(qn('w:hAnsi'), mapped_font)
                            rpr.append(rfonts)
                        else:
                            rfonts = rfonts_list[0]
                            rfonts.set(qn('w:ascii'), mapped_font)
                            rfonts.set(qn('w:eastAsia'), mapped_font)
                            rfonts.set(qn('w:hAnsi'), mapped_font)
                    except Exception as e:
                        # 如果设置样式失败，记录错误但继续
                        print(f"设置TOC样式 {style.name} 失败: {e}")
                        continue
        except Exception as e:
            # 如果整体设置失败，记录错误但继续
            print(f"设置TOC样式失败: {e}")
            # 不抛出异常，让文档生成继续
    
    def _add_bookmark_to_paragraph(self, paragraph, bookmark_name):
        """
        为段落添加书签
        """
        try:
            # 获取段落元素
            para_elem = paragraph._element
            
            # 创建唯一的书签ID（使用简单的计数器或基于名称的哈希）
            bookmark_id = str(abs(hash(bookmark_name)) % 100000)
            
            # 查找或创建书签开始标记
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), bookmark_id)
            bookmark_start.set(qn('w:name'), bookmark_name)
            
            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), bookmark_id)
            
            # 在段落开头插入书签（在第一个run之前）
            if paragraph.runs:
                first_run = paragraph.runs[0]
                first_run._element.addprevious(bookmark_start)
            else:
                # 如果没有run，创建一个空的run
                run = paragraph.add_run()
                run._element.addprevious(bookmark_start)
            
            # 在段落末尾添加书签结束标记
            para_elem.append(bookmark_end)
        except Exception as e:
            print(f"添加书签失败: {e}")
            raise
    
    def _add_pageref_field(self, run, bookmark_name):
        """
        添加PAGEREF字段来引用书签的页码
        """
        try:
            # 创建PAGEREF字段
            fld_char_begin = OxmlElement('w:fldChar')
            fld_char_begin.set(qn('w:fldCharType'), 'begin')
            
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = f'PAGEREF {bookmark_name} \\h'
            
            fld_char_sep = OxmlElement('w:fldChar')
            fld_char_sep.set(qn('w:fldCharType'), 'separate')
            
            # 添加页码占位文本
            fld_char_text = OxmlElement('w:t')
            fld_char_text.text = '1'  # 占位文本，Word会自动更新为实际页码
            
            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')
            
            # 添加到run元素
            run._element.append(fld_char_begin)
            run._element.append(instr_text)
            run._element.append(fld_char_sep)
            run._element.append(fld_char_text)
            run._element.append(fld_char_end)
        except Exception as e:
            print(f"添加PAGEREF字段失败: {e}")
            # 如果失败，添加占位文本
            run.text = "..."